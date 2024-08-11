"""
create_sample_resume_docx.py
------------------------------------------------
Author: Brian Richmond
Created on: 14 July 2024
File name: create_sample_resume_docx.py
Revised:

Description:
This script creates a sample résumé in DOCX format using the python-docx library.
The résumé contains basic details like name, email, experience, and skills.

Usage:
    Run this script to generate a sample resume DOCX file.

Example:
    python create_sample_resume_docx.py
"""

from docx import Document

# Create a Document
doc = Document()

# Add a title
doc.add_heading('John Doe', level=1)

# Add email
doc.add_paragraph('Email: john.doe@example.com')

# Add Experience
doc.add_heading('Experience:', level=2)
doc.add_paragraph('- Software Engineer at ABC Corp\n- Senior Developer at XYZ Inc.')

# Add Skills
doc.add_heading('Skills:', level=2)
doc.add_paragraph('Python, Java, SQL')

# Save the document
doc.save('sample_resume.docx')
